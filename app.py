import streamlit as st
import ollama
from pypdf import PdfReader
from docx import Document
import os
import time
import chromadb
from chromadb.utils import embedding_functions
from typing import List, Tuple

# Initialize ChromaDB
client = chromadb.PersistentClient(path="./chroma_db")
sentence_transformer_ef = embedding_functions.SentenceTransformerEmbeddingFunction(model_name="all-MiniLM-L6-v2")

try:
    collection = client.get_collection(name="documents", embedding_function=sentence_transformer_ef)
except:
    collection = client.create_collection(name="documents", embedding_function=sentence_transformer_ef)

def chunk_text(text: str, chunk_size: int = 1000) -> List[str]:
    """Split text into chunks of approximately chunk_size characters"""
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        # Try to split at sentence boundary
        if end < len(text):
            while end > start and text[end] not in {'.', '!', '?', '\n'}:
                end -= 1
            if end == start:  # No sentence boundary found
                end = start + chunk_size
        chunks.append(text[start:end].strip())
        start = end
    return chunks

def process_document(uploaded_file, progress_bar=None, status_text=None):
    """Extract text from document and store in ChromaDB with progress tracking"""
    text = ""
    
    # Update status
    if status_text:
        status_text.text(f"Extracting text from {uploaded_file.name}...")
    
    if uploaded_file.type == "application/pdf":
        reader = PdfReader(uploaded_file)
        total_pages = len(reader.pages)
        for i, page in enumerate(reader.pages):
            text += page.extract_text()
            if progress_bar:
                progress_bar.progress((i + 1) / (total_pages * 2))  # First half is for extraction
    
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(uploaded_file)
        total_paras = len(doc.paragraphs)
        for i, para in enumerate(doc.paragraphs):
            text += para.text + "\n"
            if progress_bar:
                progress_bar.progress((i + 1) / (total_paras * 2))  # First half is for extraction
    
    elif uploaded_file.type == "text/plain":
        text = str(uploaded_file.read(), "utf-8")
        if progress_bar:
            progress_bar.progress(0.5)  # Mark extraction as 50% complete
    
    # Update status
    if status_text:
        status_text.text(f"Chunking and storing {uploaded_file.name} in database...")
    
    # Split text into chunks
    chunks = chunk_text(text)
    
    # Store in ChromaDB
    ids = [f"{uploaded_file.name}-{i}" for i in range(len(chunks))]
    
    # Add chunks in batches for smoother progress updates
    batch_size = max(1, len(chunks) // 10)  # Create 10 progress updates
    for i in range(0, len(chunks), batch_size):
        end_idx = min(i + batch_size, len(chunks))
        collection.add(
            documents=chunks[i:end_idx],
            ids=ids[i:end_idx],
            metadatas=[{"source": uploaded_file.name} for _ in range(i, end_idx)]
        )
        if progress_bar:
            # Calculate progress for second half (storage)
            extraction_half = 0.5  # First 50% was for extraction
            storage_progress = (end_idx / len(chunks)) * 0.5  # Second 50% for storage
            progress_bar.progress(extraction_half + storage_progress)
    
    # Complete the progress
    if progress_bar:
        progress_bar.progress(1.0)
    if status_text:
        status_text.text(f"Completed processing {uploaded_file.name}")
    
    return len(chunks)

@st.cache_data(ttl=300)  # Cache results for 5 minutes
def retrieve_relevant_chunks(query: str, k: int = 5) -> Tuple[List[str], List[str]]:
    """Retrieve relevant document chunks from ChromaDB with caching for performance"""
    results = collection.query(
        query_texts=[query],
        n_results=k
    )
    return results['documents'][0], results['metadatas'][0]

@st.cache_data(ttl=60, show_spinner=False)  # Cache for 1 minute
def generate_response(query: str, context: str, temp: float = 0.7, model: str = "gemma3") -> str:
    """Generate response using Ollama with RAG context and caching"""
    prompt = f"""Use the following context to answer the question. If you don't know the answer, say you don't know.

Context:
{context}

Question: {query}

Answer:"""
    
    # Add optimization options
    response = ollama.generate(
        model=model,
        prompt=prompt,
        options={
            'temperature': temp,
            'num_predict': 512,  # Limit token generation for faster responses
            'num_ctx': 2048      # Set context window
        }
    )
    return response['response']

# Initialize session states
if "messages" not in st.session_state:
    st.session_state.messages = []
if "uploaded_files" not in st.session_state:
    st.session_state.uploaded_files = []

# Initialize performance tracking
if "performance_metrics" not in st.session_state:
    st.session_state.performance_metrics = {
        "total_queries": 0,
        "avg_response_time": 0,
        "last_response_time": 0
    }

# App title
st.title("📄 Document Q&A Assistant")

# Sidebar for document upload
with st.sidebar:
    st.header("Document Management")
    uploaded_files = st.file_uploader(
        "Upload documents",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True
    )
    
    st.markdown("---")
    st.header("Settings")
    temperature = st.slider("Temperature", 0.0, 1.0, 0.7)
    model_name = st.selectbox("Model", ["gemma3"], index=0)
    st.markdown("ℹ️ All processing happens locally")
    
    if uploaded_files and st.button("Process Documents", use_container_width=True):
        # Create a modal-like container for progress tracking
        progress_container = st.container()
        
        with progress_container:
            col1, col2, col3 = st.columns([1, 10, 1])
            with col2:
                st.markdown("### Processing Documents")
                progress_bar = st.progress(0)
                status_text = st.empty()
                metrics_col1, metrics_col2 = st.columns(2)
                total_chunks_metric = metrics_col1.empty()
                eta_metric = metrics_col2.empty()
                
                start_time = time.time()
                total_chunks = 0
                files_processed = 0
                
                for uploaded_file in uploaded_files:
                    if uploaded_file.name not in st.session_state.uploaded_files:
                        status_text.text(f"Starting to process {uploaded_file.name}...")
                        
                        # Process the document with progress tracking
                        chunks_count = process_document(uploaded_file, progress_bar, status_text)
                        total_chunks += chunks_count
                        files_processed += 1
                        
                        # Calculate and display metrics
                        elapsed = time.time() - start_time
                        eta = (elapsed / files_processed) * (len(uploaded_files) - files_processed) if files_processed > 0 else 0
                        total_chunks_metric.metric("Chunks Created", f"{total_chunks}")
                        eta_metric.metric("Time Remaining", f"{eta:.1f}s")
                        
                        st.session_state.uploaded_files.append(uploaded_file.name)
                
                # Final status update
                progress_bar.progress(1.0)
                status_text.text("✅ Processing completed!")
                
                # Show completion message with stats
                st.success(f"Successfully processed {files_processed} document(s) into {total_chunks} searchable chunks.")
                st.balloons()  # Add a celebratory effect
                st.markdown("### 🎉 Your documents are now ready!")
                st.markdown("You can start asking questions about your documents in the chat below.")

# Display chat messages
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# Optional: Display performance metrics in an expandable section
with st.sidebar:
    if st.session_state.performance_metrics["total_queries"] > 0:
        with st.expander("Performance Metrics"):
            st.metric("Average Response Time", f"{st.session_state.performance_metrics['avg_response_time']:.2f} seconds")
            st.metric("Last Response Time", f"{st.session_state.performance_metrics['last_response_time']:.2f} seconds")
            st.metric("Total Queries", f"{st.session_state.performance_metrics['total_queries']}")

# Chat input
if prompt := st.chat_input("Ask about your documents..."):
    # Start timing the query response
    query_start_time = time.time()
    
    # Add user message to chat history
    st.session_state.messages.append({"role": "user", "content": prompt})
    
    # Display user message
    with st.chat_message("user"):
        st.markdown(prompt)
    
    # Display assistant response
    with st.chat_message("assistant"):
        message_placeholder = st.empty()
        full_response = ""
        
        # Add loading indicator
        with st.status("Searching documents for relevant information...", expanded=True) as status:
            # Retrieve relevant context with improved k value
            st.write("🔍 Finding relevant information...")
            chunks, metadata = retrieve_relevant_chunks(prompt)
            context = "\n\n".join(chunks)
            
            # Show found sources
            sources = list(set([m['source'] for m in metadata]))
            st.write(f"📚 Found information in {len(sources)} document(s)")
            
            # Generate response with context
            st.write("💭 Generating response...")
            response = generate_response(prompt, context, temp=temperature, model=model_name)
            status.update(label="✅ Answer ready!", state="complete", expanded=False)
        
        # Stream the response more efficiently with batch updates
        words = response.split()
        total_words = len(words)
        update_frequency = max(1, total_words // 20)  # Update about 20 times total
        
        for i in range(0, total_words, update_frequency):
            end_idx = min(i + update_frequency, total_words)
            full_response += " ".join(words[i:end_idx]) + " "
            message_placeholder.markdown(full_response + "▌")
            # Use a much shorter sleep time for faster updates
            time.sleep(0.01)
        
        # Add sources
        sources = list(set([m['source'] for m in metadata]))
        if sources:
            full_response += f"\n\nSources: {', '.join(sources)}"
        
        message_placeholder.markdown(full_response)
    
    # Add assistant response to chat history
    st.session_state.messages.append({"role": "assistant", "content": full_response})
    
    # Update performance metrics
    end_time = time.time()
    query_time = end_time - query_start_time
    
    # Update metrics
    st.session_state.performance_metrics["total_queries"] += 1
    st.session_state.performance_metrics["last_response_time"] = query_time
    
    # Calculate running average
    prev_avg = st.session_state.performance_metrics["avg_response_time"] 
    prev_count = st.session_state.performance_metrics["total_queries"] - 1
    
    if prev_count > 0:
        st.session_state.performance_metrics["avg_response_time"] = (prev_avg * prev_count + query_time) / st.session_state.performance_metrics["total_queries"]
    else:
        st.session_state.performance_metrics["avg_response_time"] = query_time
